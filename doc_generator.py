import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Colors for professional design
COLOR_PRIMARY = RGBColor(31, 78, 121)    # Deep Steel Blue
COLOR_SECONDARY = RGBColor(128, 128, 128) # Slate Gray
COLOR_SUCCESS = RGBColor(40, 167, 69)    # Forest Green
COLOR_FAIL = RGBColor(220, 53, 69)       # Crimson Red
COLOR_TEXT = RGBColor(51, 51, 51)        # Charcoal

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner margins for table cells."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx_report(summary_text, timeline, checklist, eligibility_results, vendor_profile) -> io.BytesIO:
    """
    Generates a structured, styled Word document (DOCX) for downloading.
    """
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("TENDER SIMPLIFIER REPORT")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Prepared for: {vendor_profile.get('name', 'Valued Vendor')} | Generated automatically via Tender Simplifier")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_SECONDARY
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(12)

    # Helper function to add styled headings
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        # Add bottom border/line using paragraph border in XML if needed, or simple horizontal line
        doc.add_paragraph("__________________________________________________________________").paragraph_format.space_after = Pt(12)

    # 2. Executive Summary Section
    add_section_heading("1. Executive Summary & Scope")
    
    # Standardize output text formatting
    summary_p = doc.add_paragraph()
    summary_run = summary_p.add_run(summary_text)
    summary_run.font.name = 'Arial'
    summary_run.font.size = Pt(10.5)
    summary_run.font.color.rgb = COLOR_TEXT

    # 3. Eligibility Check
    add_section_heading("2. Vendor Eligibility Audit")
    
    eligibility_p = doc.add_paragraph()
    overall_status = "ELIGIBLE" if eligibility_results.get("eligible", False) else "NOT ELIGIBLE"
    status_color = COLOR_SUCCESS if overall_status == "ELIGIBLE" else COLOR_FAIL
    
    eligibility_p.add_run("Overall Assessment: ").font.bold = True
    status_run = eligibility_p.add_run(overall_status)
    status_run.font.bold = True
    status_run.font.color.rgb = status_color
    status_run.font.size = Pt(12)
    
    # Table for detailed eligibility rules
    elig_table = doc.add_table(rows=1, cols=4)
    elig_table.style = 'Light Shading Accent 1'
    
    # Header Row
    hdr_cells = elig_table.rows[0].cells
    hdr_cells[0].text = 'Criterion'
    hdr_cells[1].text = 'Tender Requirement'
    hdr_cells[2].text = 'Vendor Value'
    hdr_cells[3].text = 'Status'
    
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79") # COLOR_PRIMARY in Hex
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = 'Arial'
                r.font.size = Pt(10)
    
    # Populate Table Rows
    criteria_keys = [
        ("Turnover Check", "turnover"),
        ("Experience Check", "experience"),
        ("Certifications Check", "certifications")
    ]
    
    for label, key in criteria_keys:
        res = eligibility_results.get(key, {})
        row_cells = elig_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(res.get("required", "N/A"))
        row_cells[2].text = str(res.get("vendor", "N/A"))
        
        status = res.get("status", "N/A")
        status_cell = row_cells[3]
        status_cell.text = status
        
        # Color-code status cell
        if status == "PASS":
            set_cell_background(status_cell, "D4EDDA") # Light Green Hex
            for p in status_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(21, 87, 36)
                    r.font.bold = True
        elif status == "FAIL":
            set_cell_background(status_cell, "F8D7DA") # Light Red Hex
            for p in status_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(114, 28, 36)
                    r.font.bold = True
                    
        for cell in row_cells:
            set_cell_margins(cell, 80, 80, 100, 100)
            for p in cell.paragraphs:
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(9.5)
                
    doc.add_paragraph("").paragraph_format.space_after = Pt(12)

    # 4. Critical Timeline & Dates
    add_section_heading("3. Critical Timeline & Key Dates")
    
    if timeline:
        time_table = doc.add_table(rows=1, cols=3)
        time_table.style = 'Light Shading Accent 1'
        
        hdr_cells = time_table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Milestone Event'
        hdr_cells[2].text = 'Description'
        
        for cell in hdr_cells:
            set_cell_background(cell, "1F4E79")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
                    
        for milestone in timeline:
            row_cells = time_table.add_row().cells
            row_cells[0].text = str(milestone.get("date", "N/A"))
            row_cells[1].text = str(milestone.get("milestone", "N/A"))
            row_cells[2].text = str(milestone.get("description", "N/A"))
            
            for cell in row_cells:
                set_cell_margins(cell, 80, 80, 100, 100)
                for p in cell.paragraphs:
                    p.style.font.name = 'Arial'
                    p.style.font.size = Pt(9.5)
    else:
        doc.add_paragraph("No specific milestones found in the timeline extraction.")
        
    doc.add_paragraph("").paragraph_format.space_after = Pt(12)

    # 5. Required Documents Compliance Checklist
    add_section_heading("4. Required Documents Checklist")
    
    if checklist:
        check_table = doc.add_table(rows=1, cols=5)
        check_table.style = 'Light Shading Accent 1'
        
        hdr_cells = check_table.rows[0].cells
        hdr_cells[0].text = 'Status'
        hdr_cells[1].text = 'Document Name'
        hdr_cells[2].text = 'Category'
        hdr_cells[3].text = 'Mandatory'
        hdr_cells[4].text = 'Notes / Details'
        
        for cell in hdr_cells:
            set_cell_background(cell, "1F4E79")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
                    
        for doc_item in checklist:
            row_cells = check_table.add_row().cells
            row_cells[0].text = "[  ] Pending" # Checkbox column
            row_cells[1].text = str(doc_item.get("document_name", "N/A"))
            row_cells[2].text = str(doc_item.get("category", "General"))
            
            is_mandatory = doc_item.get("mandatory", True)
            row_cells[3].text = "Yes" if is_mandatory else "No"
            if is_mandatory:
                for p in row_cells[3].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = COLOR_FAIL
            
            row_cells[4].text = str(doc_item.get("description", "N/A"))
            
            for cell in row_cells:
                set_cell_margins(cell, 80, 80, 100, 100)
                for p in cell.paragraphs:
                    p.style.font.name = 'Arial'
                    p.style.font.size = Pt(9.5)
    else:
        doc.add_paragraph("No specific documents extracted in the checklist.")

    # Save to BytesIO
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file
